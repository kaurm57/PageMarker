from docx import Document 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

def run_clear(x):
    """Reset formatting for a run while keeping the text."""
    x.font.bold = False
    x.font.italic = False
    x.font.underline = False
    x.font.strike = False
    x.font.all_caps = False
    x.font.small_caps = False
    x.font.subscript = False
    x.font.superscript = False
    x.font.size = Pt(12)
    x.font.color.rgb = RGBColor(0, 0, 0)
    x.font.highlight_color = None
    x.font.shadow = None
    x.font.outline = None
    x.font.rtl = False
    x.font.complex_script = False
    x.font.name = 'Verdana'

def para_clear(para_format):
    """Reset paragraph formatting while keeping the paragraph style."""
    para_format.left_indent = None
    para_format.right_indent = None
    para_format.first_line_indent = None
    para_format.space_before = None
    para_format.space_after = None
    para_format.line_spacing = None
    para_format.border_top = None
    para_format.border_bottom = None
    para_format.border_left = None
    para_format.border_right = None

def clear_formatting(doc_path, edited_doc):
    """Clear all paragraph and run formatting in the document."""
    doc = Document(doc_path)

    for para in doc.paragraphs:
        para_clear(para.paragraph_format)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in para.runs:
            run_clear(run)

    doc.save(edited_doc)

def format_text(doc_path, edited_doc):
    """Apply custom formatting based on the paragraph style."""
    doc = Document(doc_path)

    for para in doc.paragraphs:
        para_format = para.paragraph_format
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        style_name = para.style.name

        if style_name == "Title":
            para_format.space_before = Pt(6)
            para_format.space_after = Pt(15)
            for run in para.runs:
                run.font.name = "Verdana"
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif style_name == "Heading 1":
            para_format.space_before = Pt(24)
            para_format.space_after = Pt(9)
            for run in para.runs:
                run.font.name = "Verdana"
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif style_name == "Heading 2":
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(9)
            for run in para.runs:
                run.font.name = "Verdana"
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif style_name.startswith("Heading") and style_name != "Heading 6":
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(6)
            for run in para.runs:
                run.font.name = "Verdana"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif style_name == "Heading 6":
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(10)
            for run in para.runs:
                run.font.name = "Verdana"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(10)
            for run in para.runs:
                run.font.name = "Verdana"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(edited_doc)

def remove_style_borders(doc_path, edited_path, style_name_to_fix):
    """Remove paragraph border from the specified style."""
    doc = Document(doc_path)
    style_to_fix = doc.styles[style_name_to_fix]
    style_element = style_to_fix._element

    pPr = style_element.find(qn('w:pPr'))
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)

    doc.save(edited_path)
