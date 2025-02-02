from docx import Document 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import re

def add_page_num_and_make_heading(doc_path, editted_doc, page_count):
    # Open the document
    doc = Document(doc_path)
    new_doc = Document()

    for para in doc.paragraphs:
        for run in para.runs:
            for br in run._element.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    page_count += 1
                    run._element.remove(br)
                    run.add_text(f"[Page {page_count}]")  # No newline space

        # After removing page breaks and adding page numbers, format page numbers as headings
        match = re.search(r"\[Page \d+\]", para.text)
        if match:
            # Add an empty paragraph before the page number
            new_doc.add_paragraph()  # Empty paragraph to add space before the page number
            page_text = match.group(0)
            page_heading = new_doc.add_paragraph(page_text)
            page_heading.style = "Heading 6"
        else:
            new_para = new_doc.add_paragraph(para.text)
            new_para.style = para.style

    # Save the final result
    new_doc.save(editted_doc)

# Call the function with paths to your document
add_page_num_and_make_heading("./testing.docx", "./final-result.docx", 0)


def run_clear(x):
    """Reset all formatting for a run while keeping the text."""
    # Reset font properties
    x.font.bold = False
    x.font.italic = False
    x.font.underline = False
    x.font.strike = False
    x.font.all_caps = False
    x.font.small_caps = False
    x.font.subscript = False
    x.font.superscript = False
    x.font.size = Pt(12)  # Default size (you can adjust)
    x.font.color.rgb = RGBColor(0, 0, 0)  # Reset font color to black
    x.font.highlight_color = None
    x.font.shadow = None
    x.font.outline = None
    x.font.rtl = False
    x.font.complex_script = False
    x.font.name = 'Verdana'  # Default font

def para_clear(para_format):
    """Reset paragraph formatting while keeping the paragraph style."""
    para_format.left_indent = None
    para_format.right_indent = None
    para_format.first_line_indent = None
    para_format.space_before = None
    para_format.space_after = None
    para_format.line_spacing = None
    para_format.border_top = None  # Clear top border
    para_format.border_bottom = None  # Clear bottom border
    para_format.border_left = None  # Clear left border
    para_format.border_right = None  # Clear right border

def clear_formatting(doc_path, edited_doc):
    """Clear all formatting in the Word document."""
    doc = Document(doc_path)

    for para in doc.paragraphs:
        para_format = para.paragraph_format
        para_clear(para_format)  # Reset paragraph formatting
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Default all to left-aligned
        
        for run in para.runs:
            run_clear(run)  # Apply run_clear to remove all formatting

    # Save the edited document
    doc.save(edited_doc)

# Run the function
clear_formatting("./final-result.docx", "./final-final-result.docx")

def format_text(doc_path, edited_doc):
    doc = Document(doc_path)

    for para in doc.paragraphs:
        para_format = para.paragraph_format
        #para_clear(para_format)  # Reset paragraph formatting
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Default all to left-aligned
        
        # Apply styles based on paragraph style name
        style_name = para.style.name
        
        # Title
        if style_name == "Title":
            para_format.space_before = Pt(6)
            para_format.space_after = Pt(15)
            for run in para.runs:
                #run_clear(run)
                run.font.name = "Verdana"
                run.font.size = Pt(16)
                run.font.bold = True  # Only Title is bold by default
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Heading 1
        elif style_name == "Heading 1":
            para_format.space_before = Pt(24)
            para_format.space_after = Pt(9)
            for run in para.runs:
                #run_clear(run)
                run.font.name = "Verdana"
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Heading 2
        elif style_name == "Heading 2":
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(9)
            for run in para.runs:
                #run_clear(run)
                run.font.name = "Verdana"
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Heading 3 and above (excluding Heading 6)
        elif style_name.startswith("Heading") and style_name != "Heading 6":
            para_format.space_before = Pt(18)
            para_format.space_after = Pt(6)
            for run in para.runs:
                #run_clear(run)
                run.font.name = "Verdana"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Heading 6
        elif style_name == "Heading 6":
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(10)
            for run in para.runs:
                #run_clear(run)
                run.font.name = "Verdana"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Normal text
        else:
            para_format.space_before = Pt(0)
            para_format.space_after = Pt(10)
            for run in para.runs:
                #run_clear(run)
                run.font.name = "Verdana"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black

    # Save the formatted document
    doc.save(edited_doc)

format_text("./final-final-result.docx", "./full-and-final-result.docx")