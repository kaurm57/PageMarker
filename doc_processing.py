from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
import re
import io

def add_page_num(doc_path, edited_doc, page_count):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Inches
    import io, re

    doc = Document(doc_path)
    new_doc = Document()

    for para in doc.paragraphs:
        marker_found = False
        # Process runs to detect page breaks
        for run in para.runs:
            # Look for all page break elements
            for br in run._element.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    page_count += 1
                    run._element.remove(br)
                    marker_found = True

        # Add the current paragraph's content to the new document
        new_para = new_doc.add_paragraph()
        new_para.style = para.style
        for run in para.runs:
            if run._element.xpath('.//w:drawing'):
                # Handle images
                alt_text = ""
                docPr_elements = run._element.xpath('.//wp:docPr')
                if docPr_elements:
                    alt_text = docPr_elements[0].get("descr", "")
                blip_elements = run._element.xpath('.//a:blip')
                if blip_elements:
                    rId = blip_elements[0].get(qn('r:embed'))
                    image_part = doc.part.related_parts[rId]
                    image_stream = io.BytesIO(image_part.blob)
                    shape = new_doc.add_picture(image_stream, width=Inches(4))
                    docPr_new = shape._inline.xpath('./wp:docPr')
                    if docPr_new:
                        docPr_new[0].set("descr", alt_text)
            else:
                new_para.add_run(run.text)

        # If a page break was encountered in this paragraph, add a single marker heading
        if marker_found:
            page_heading = new_doc.add_paragraph(f"[Page {page_count}]")
            page_heading.style = "Heading 6"
            new_doc.add_paragraph()  # Optional extra spacing

    new_doc.save(edited_doc)
    return page_count

